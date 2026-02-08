from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Create a new Document
doc = Document()

# Set default font to Times New Roman (standard for academic reports)
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Function to add formatted paragraph
def add_para(text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

# --- PAGE 1: COVER PAGE ---
add_para("Project Group No: ______", bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
add_para("", space_after=12) # Spacing

add_para("BOL BHARAT — AI-Powered Civic Issue Reporting Platform", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

add_para("Project Domain: Internet of Things / Web Development", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para("SY Community Engineering Project Synopsis", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Bachelor of Engineering", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("in", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Computer Engineering", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para("By", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
# Group Members (Placeholders from MD)
members = [
    "Kartik [Roll No: XX]",
    "[Member 2 Name] [Roll No: XX]",
    "[Member 3 Name] [Roll No: XX]",
    "[Member 4 Name] [Roll No: XX]"
]
for m in members:
    add_para(m, bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

add_para("", space_after=18)

add_para("Supervisor Name", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("[Your Guide's Name]", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)

add_para("Department of Computer Engineering,", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Vidya Pratishthan's Kamalnayan Bajaj Institute of Engineering & Technology,", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Baramati. Dist: Pune, Maharashtra", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Academic Year 2025-26", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# --- PAGE 2: DECLARATION ---
add_para("We,", bold=True, size=12)
add_para("Project Approval Declaration", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

# Member List for Declaration
for i, m in enumerate(members, 1):
    add_para(f"{i}. {m} [PRN No: __________]", size=12)

add_para("", space_after=12)

declaration_text = (
    "hereby declare that we are submitting the curricular project synopsis entitled "
    "\"BOL BHARAT — AI-Powered Civic Issue Reporting Platform\", under the broad domain of Internet of Things (IoT).\n"
    "This project has been carried out with the due approval of our Project Guide and the Head of the "
    "Department, and is submitted in partial fulfilment of the requirements for the award of the degree "
    "of Bachelor of Engineering during the academic year 2025-26."
)
p = add_para(declaration_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
p.paragraph_format.first_line_indent = Inches(0.5)

add_para("", space_after=36)

# Signatures
table = doc.add_table(rows=1, cols=2)
table.autofit = True
row = table.rows[0]
c1 = row.cells[0].paragraphs[0]
c1.add_run("[Your Guide's Name]\nGuide\nDepartment of Computer Engineering").bold = True
c1.alignment = WD_ALIGN_PARAGRAPH.LEFT

c2 = row.cells[1].paragraphs[0]
c2.add_run("Dr. Arvind Jagtap\nHead\nDepartment of Computer Engineering").bold = True
c2.alignment = WD_ALIGN_PARAGRAPH.RIGHT

add_para("", space_after=24)
add_para("Place: Baramati", bold=True)
add_para("Date: ____________", bold=True)

doc.add_page_break()

# --- PAGE 3: ABSTRACT ---
add_para("Abstract", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

abstract_text = (
    "In India, citizens face significant challenges in reporting civic issues like potholes, "
    "water leakage, garbage accumulation, and streetlight failures to the appropriate government departments. "
    "Current systems are fragmented, non-transparent, and manual. "
    "BOL BHARAT (Speak Up, India!) is an AI-powered civic issue reporting platform that empowers citizens to "
    "report issues with photos and location. The system utilizes Gemini AI and local Machine Learning "
    "algorithms to auto-categorize complaints and route them to the correct government department based on pincode. "
    "Key features include real-time tracking via Firebase, community engagement through upvoting, and a "
    "mobile-responsive design."
)
add_para(abstract_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

doc.add_page_break()

# --- PAGE 4: INTRODUCTION ---
add_para("1. Introduction", bold=True, size=14, space_after=12)

intro_text = (
    "Civic engagement is crucial for a functioning democracy, yet reporting infrastructure failures "
    "remains difficult. Existing complaint systems are often fragmented across different departments, "
    "manual, and lack transparency regarding complaint status.\n\n"
    "BOL BHARAT addresses these issues by providing a centralized, AI-driven platform. "
    "The proposed solution allows citizens to report issues via a wizard interface. "
    "Using AI (Gemini and Local ML), the system automates the categorization and description "
    "of issues. Furthermore, it implements 'Smart Routing' to direct complaints to the correct "
    "department based on the pincode, and includes a community upvoting system to highlight urgent issues."
)
add_para(intro_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

add_para("", space_after=24)
add_para("Fig: Block Diagram", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
# Placeholder for image
add_para("[Paste System Architecture Diagram from MD Section 5.1 Here]", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

doc.add_page_break()

# --- PAGE 5: LITERATURE REVIEW ---
add_para("2. Background of the Invention/Project (Literature Review)", bold=True, size=14, space_after=12)

# Converting MD Table to Narrative
lit_reviews = [
    ("SWACHHATA App (Government)", "Ministry of Housing", "It provides a centralized complaint system but is limited to a single category and lacks AI/tracking."),
    ("FixMyStreet (UK)", "mySociety", "It offers location-based reporting but lacks ML integration for auto-categorization."),
    ("SeeClickFix (USA)", "SeeClickFix Inc.", "It demonstrates strong community engagement but uses a paid model and is US-centric."),
    ("Smart City Complaint Systems", "Various Research Papers", "These highlight IoT potential but often involve high infrastructure costs."),
    ("Image Classification for Urban Issues", "IEEE/ACM Papers", "These propose CNN-based categorization but require large training datasets.")
]

for i, (title, author, finding) in enumerate(lit_reviews, 1):
    p = add_para(f"In the system/journal regarding \"{title}\" by {author}, it was observed that: {finding} [{i}].", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# --- PAGE 6: PROBLEM & OBJECTIVES ---
add_para("3. Problem Statement & Proposed Project", bold=True, size=14, space_after=12)

prob_text = (
    "In India, citizens struggle to report civic issues due to fragmented portals and a lack of transparency. "
    "The process is often manual and slow. In this project, it is expected that we implement 'BOL BHARAT', "
    "an AI-powered platform. It will use Gemini AI to auto-fill reports from images/descriptions and "
    "automatically route them to the correct department using pincode mapping."
)
add_para(prob_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24)

add_para("4. Significance & Objectives", bold=True, size=14, space_after=12)

objectives = [
    "Simplify Reporting: Create an intuitive wizard for citizens.",
    "Leverage AI: Use Gemini AI and Local ML for auto-filling and categorization.",
    "Smart Routing: Automatically direct complaints based on pincode.",
    "Community Engagement: Enable upvoting to highlight urgent issues.",
    "Real-time Tracking: Provide live status updates using Firebase.",
    "Accessibility: Ensure mobile-responsive design."
]

add_para("Objectives:", bold=True)
for i, obj in enumerate(objectives, 1):
    add_para(f"{i}. {obj}")

doc.add_page_break()

# --- PAGE 7: METHODOLOGY ---
add_para("5. Methodology (with flowchart)", bold=True, size=14, space_after=12)

method_text = (
    "The system is built on a React 18 frontend with a Firebase backend. "
    "1. Requirement Analysis: Identified gaps in current systems.\n"
    "2. System Architecture: The frontend interacts with a Service Layer that calls Gemini AI for text analysis "
    "and Local ML for image categorization.\n"
    "3. Routing Logic: A dedicated algorithm maps 6-digit pincodes to specific state departments.\n"
    "4. Community Logic: Issues are sorted by an upvoting algorithm to determine trending status."
)
add_para(method_text, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
add_para("[Refer to 5.1 System Architecture in MD for Flowchart]", align=WD_ALIGN_PARAGRAPH.CENTER, size=10)

doc.add_page_break()

# --- PAGE 8: TOOLS & PLAN ---
add_para("6. Development Tools", bold=True, size=14, space_after=12)

add_para("A. Hardware", bold=True)
add_para("1. Developer Laptops\n2. Smartphones for testing")

add_para("B. Software", bold=True)
add_para("1. Frontend: React 18, TypeScript, Tailwind CSS\n2. Backend: Firebase Realtime DB\n3. AI: Google Gemini 2.5 Flash API")

add_para("", space_after=12)
add_para("7. Project Plan", bold=True, size=14, space_after=12)
add_para("The following table outlines the 12-week development plan:", space_after=6)

# Plan Table
plan_data = [
    ("Task/Milestone", "Start Date", "End Date", "Status"),
    ("Requirement Analysis", "01-Jan-2026", "15-Jan-2026", "Completed"),
    ("Literature Survey", "16-Jan-2026", "31-Jan-2026", "Completed"),
    ("UI/UX Design", "01-Feb-2026", "07-Feb-2026", "Pending"),
    ("Frontend Dev (React)", "08-Feb-2026", "21-Feb-2026", "Pending"),
    ("Firebase Integration", "22-Feb-2026", "28-Feb-2026", "Pending"),
    ("AI Service Integration", "01-Mar-2026", "14-Mar-2026", "Pending"),
    ("Testing & Final Report", "15-Mar-2026", "31-Mar-2026", "Pending")
]

table = doc.add_table(rows=len(plan_data), cols=4)
table.style = 'Table Grid'
for i, row_data in enumerate(plan_data):
    row = table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text

add_para("", space_after=24)
add_para("References", bold=True, size=14)

refs = [
    "React Documentation — https://react.dev/",
    "Firebase Documentation — https://firebase.google.com/docs",
    "Google Gemini API — https://ai.google.dev/gemini-api/docs",
    "SWACHHATA App — https://swachhata.meitylabs.in/",
    "FixMyStreet — https://www.fixmystreet.com/"
]

for i, r in enumerate(refs, 1):
    add_para(f"[{i}] {r}")

# Save
file_path = "CEP_Synopsis_Bol_Bharat.docx"
doc.save(file_path)
print(f"Document saved to: {file_path}")